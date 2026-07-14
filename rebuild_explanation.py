"""
Rebuild code_explanation.docx to match the new report structure
with detailed step-by-step explanation following the sample assignment format.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(text, bold=False, italic=False, alignment=None, size=12):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_empty_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    add_empty_line()

add_para('Code & Solution Explanation Document', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=18)
add_empty_line()
add_para('COM7023 - Mathematics for Data Science', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_empty_line()
add_para('Urban Resource Intelligence Portfolio', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_empty_line()
add_empty_line()
add_para('Session 1: Tasks 1 and 2', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_line()
add_para(
    'This document provides a plain-English, step-by-step explanation of every Python script '
    'and every mathematical solution step. It serves as the author\'s study and defence notes - '
    'the author must be able to explain every part of this work verbally if asked.',
    italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11
)

doc.add_page_break()

# ============================================================
# TASK 1: STEADY-STATE RESOURCE DEMAND (LINEAR ALGEBRA)
# ============================================================
add_heading_styled('Task 1: Steady-State Resource Demand (Linear Algebra)', level=1)

# 1.1 What the Problem Is Asking
add_heading_styled('1.1 What the Problem Is Asking', level=2)

add_para(
    'We have four urban zones. Each zone has a certain daily resource demand (x₁, x₂, x₃, x₄). '
    'The zones are connected - demand in one zone affects the others. The problem gives us four '
    'equations that describe these relationships, and we need to find the one set of demand values '
    'that satisfies all four equations simultaneously. This is called the "steady-state" or '
    '"equilibrium" solution - the point where the system balances.'
)

add_para(
    'The system of equations describes how resource demand flows between zones. For example, '
    'the first equation (2x₁ − x₂ + x₃ = 120) tells us that Zone 1\'s own demand coefficient is 2, '
    'it loses some demand to Zone 2 (coefficient −1), gains from Zone 3 (coefficient +1), and the '
    'total net demand target is 120 units.'
)

add_para('The system of equations is:')
add_bullet('2x₁ − x₂ + x₃ = 120')
add_bullet('−x₁ + 3x₂ − x₄ = 150')
add_bullet('x₁ + x₂ + 2x₃ − x₄ = 180')
add_bullet('−x₂ + x₃ + 2x₄ = 140')

# 1.2 The Mathematics Behind It
add_heading_styled('1.2 The Mathematics Behind It - Matrix Algebra', level=2)

add_para(
    'We write the system as Ax = b, where A is the coefficient matrix (containing the numbers '
    'that multiply x₁, x₂, x₃, x₄), x is the vector of unknowns we want to find, and b is the '
    'vector of target demands (120, 150, 180, 140).'
)

add_para(
    'In matrix form: A is a 4×4 matrix, x and b are 4-element column vectors. The equation Ax = b '
    'represents the entire system compactly. Solving Ax = b means finding the vector x such that '
    'when multiplied by A, we get b.'
)

# 1.3 Verification of Matrix Properties
add_heading_styled('1.3 Why We Check Matrix Properties', level=2)

add_para('Before solving, we check three important properties of matrix A:')

add_para('1. Determinant (det(A))', bold=True)
add_para(
    'The determinant is a scalar value that tells us whether a matrix is invertible. If det(A) ≠ 0, '
    'the matrix is invertible and there is exactly one unique solution. If det(A) = 0, the system '
    'has either no solution or infinitely many solutions. For our system, det(A) = 12 ≠ 0, so a '
    'unique solution exists.'
)

add_para('2. Rank', bold=True)
add_para(
    'The rank of a matrix is the number of linearly independent rows (or columns). A full rank '
    '(rank = number of equations) means all equations provide unique information - none is a '
    'combination of others. Our matrix has rank = 4, meaning all four equations are independent.'
)

add_para('3. Condition Number', bold=True)
add_para(
    'The condition number measures how sensitive the solution is to small changes in the input. '
    'A small condition number (< 100) means the system is "well-conditioned" - small rounding '
    'errors in the coefficients will not cause large errors in the solution. Our condition number '
    'is 7.92, which is excellent.'
)

# 1.4 Step-by-Step Solution
add_heading_styled('1.4 Step-by-Step Solution Using Gaussian Elimination', level=2)

add_heading_styled('Step 1: Form the Augmented Matrix', level=3)
add_para(
    'The augmented matrix [A|b] combines A and b into a single 4×5 matrix. This is the starting '
    'point for Gaussian elimination. The vertical bar separates the coefficients from the constants.'
)
add_para(
    '[A|b] places the coefficients of each equation on the left and the constants on the right. '
    'Row 1: [2, -1, 1, 0 | 120] represents 2x₁ − x₂ + x₃ + 0x₄ = 120.'
)

add_heading_styled('Step 2: Forward Elimination (Creating Upper Triangular Form)', level=3)
add_para(
    'Gaussian elimination uses three types of row operations: (1) swapping rows, (2) multiplying '
    'a row by a non-zero constant, and (3) adding a multiple of one row to another. The goal is to '
    'create zeros below the diagonal, transforming the matrix into row-echelon form (upper triangular).'
)

add_para('Elimination of x₁ from rows 2 and 3:', bold=True)
add_para(
    'We use row 1 as the pivot row. To eliminate x₁ from row 2, we compute R₂ + 0.5R₁ because '
    'the coefficient in row 2, column 1 is -1, and we want to make it 0. The multiplier is '
    '-(−1)/2 = 0.5. Similarly, to eliminate x₁ from row 3, we compute R₃ - 0.5R₁.'
)

add_para('Elimination of x₂ from rows 3 and 4:', bold=True)
add_para(
    'Now using row 2 as the pivot row. The multiplier for row 3 is 1.5/2.5 = 0.6, so R₃ - 0.6R₂. '
    'For row 4, the multiplier is -(-1)/2.5 = 0.4, so R₄ + 0.4R₂.'
)

add_para('Elimination of x₃ from row 4:', bold=True)
add_para(
    'Using row 3 as the pivot row. Both row 3 and row 4 have 1.2 in the x₃ column, so R₄ - R₃ '
    'eliminates x₃ from row 4.'
)

add_heading_styled('Step 3: Back Substitution', level=3)
add_para(
    'Once the matrix is in upper triangular form, we solve from the bottom up. The last equation '
    '(row 4) has only one unknown (x₄), so we solve for x₄ first. Then row 3 gives us x₃ using '
    'the known x₄ value. We continue upward until all unknowns are found.'
)

add_para('The solution is found by working backwards:')
add_bullet('From row 4: 2x₄ = 230 → x₄ = 115')
add_bullet('From row 3: 1.2x₃ − 0.4(115) = −6 → x₃ = 33.33')
add_bullet('From row 2: 2.5x₂ + 0.5(33.33) − 115 = 210 → x₂ = 123.33')
add_bullet('From row 1: 2x₁ − 123.33 + 33.33 = 120 → x₁ = 105')

add_para('Final solution:', bold=True)
add_bullet('x₁ = 105.00 units/day (Zone 1)')
add_bullet('x₂ = 123.33 units/day (Zone 2)')
add_bullet('x₃ = 33.33 units/day (Zone 3)')
add_bullet('x₄ = 115.00 units/day (Zone 4)')

# 1.5 Python Code - Line by Line
add_heading_styled('1.5 The Python Code - Line by Line', level=2)

add_para(
    'The script is in code/task1_linear_algebra.py. Here is what each section does:'
)

add_para('Section 1 - Setup (Lines 14-16):', bold=True)
add_para(
    'We import numpy (as np) for matrix operations, seaborn (as sns) for plotting, and matplotlib '
    '(as plt) for figure control. We define matrix A as a 4×4 NumPy array and vector b as a 1D array. '
    'The "dtype=float" ensures we use decimal numbers, not integers.'
)

add_para('Section 2 - Existence and Uniqueness Checks (Lines 46-74):', bold=True)
add_para(
    'np.linalg.det(A) computes the determinant. If it is not zero, the matrix is invertible and '
    'there is exactly one solution. np.linalg.matrix_rank(A) checks if all rows are independent. '
    'np.linalg.cond(A) computes the condition number - a measure of how sensitive the solution is '
    'to small changes in the input. A condition number of 7.92 (as we got) is very good.'
)

add_para('Section 3 - Solving (Lines 84-96):', bold=True)
add_para(
    'np.linalg.solve(A, b) is the main solver. It uses LAPACK routines under the hood - these are '
    'highly optimised and numerically stable. We then verify by computing A @ x (matrix-vector '
    'multiplication) and comparing to b. The residual (maximum difference) should be near zero - '
    'we got 2.84 × 10⁻¹⁴, which is essentially zero (machine precision).'
)

add_para('Section 4 - Visualisation (Lines 122-152):', bold=True)
add_para(
    'We use seaborn\'s barplot to create a clean bar chart. Each bar represents one zone\'s demand. '
    'We add value labels on top, a horizontal dashed line for the average, and proper axis labels '
    'with units. The figure is saved at 150 dpi as required by the brief.'
)

# 1.6 Understanding the Results
add_heading_styled('1.6 Understanding the Results', level=2)

add_para(
    'The solution tells us: Zone 1 needs 105 units/day, Zone 2 needs 123.33, Zone 3 needs 33.33, '
    'and Zone 4 needs 115.00. Zone 2 has the highest demand (about 3.7 times Zone 3).'
)

add_para(
    'The fact that det(A) = 12 (non-zero) and rank = 4 (full) means this is the ONLY possible '
    'steady state - the system will always converge to these values. This gives planners certainty: '
    'they can allocate resources based on these numbers and know the system will balance.'
)

add_para(
    'Urban planning implications: Zone 2 likely represents a high-density commercial or residential '
    'area requiring the most resources. Zone 3, with the lowest demand at only 33.33 units/day, could '
    'be a green zone, park, or low-density residential area. The unique solution means resource '
    'distribution has a single predictable outcome, enabling confident planning.'
)

doc.add_page_break()

# ============================================================
# TASK 2: DEMAND OVER TIME (CALCULUS)
# ============================================================
add_heading_styled('Task 2: Demand Over Time (Calculus)', level=1)

# 2.1 What the Problem Is Asking
add_heading_styled('2.1 What the Problem Is Asking', level=2)

add_para(
    'We have a function D(t) = 4t³ − 18t² + 24t + 90 that describes how resource demand changes '
    'over a 4-hour period. We need to find:'
)

add_bullet('When demand peaks and troughs (the critical points).')
add_bullet('Whether those peaks are maxima or minima (classification using second derivative).')
add_bullet('Where the rate of change itself changes (the inflection point).')
add_bullet('When demand is increasing vs decreasing (intervals of monotonicity).')
add_bullet('Practical planning implications for resource management.')

# 2.2 The Mathematics Behind It
add_heading_styled('2.2 The Mathematics Behind It - Differential Calculus', level=2)

add_para(
    'We use differential calculus, which is the study of rates of change. The key tools are:'
)

add_para('First Derivative D\'(t):', bold=True)
add_para(
    'The first derivative tells us the instantaneous rate of change of demand at any time t. '
    'If D\'(t) > 0, demand is increasing. If D\'(t) < 0, demand is decreasing. Where D\'(t) = 0, '
    'demand is momentarily constant - these are candidate points for maxima or minima.'
)

add_para('Second Derivative D\'\'(t):', bold=True)
add_para(
    'The second derivative tells us about the concavity of the function - whether the curve is '
    'bending upward (like a cup, D\'\'(t) > 0) or downward (like a cap, D\'\'(t) < 0). More importantly, '
    'it helps classify critical points: if D\'\'(t) > 0 at a critical point, it is a local minimum; '
    'if D\'\'(t) < 0, it is a local maximum.'
)

add_para('Inflection Point:', bold=True)
add_para(
    'Where D\'\'(t) = 0, the concavity changes. This is called an inflection point. Even though '
    'demand may still be decreasing at this point, the rate of decrease is slowing down, signalling '
    'a forthcoming change in direction.'
)

# 2.3 Step-by-Step Solution
add_heading_styled('2.3 Step-by-Step Solution', level=2)

add_heading_styled('Step 1: Apply the Power Rule', level=3)
add_para(
    'The power rule states that for any term atⁿ, the derivative is n·a·tⁿ⁻¹. Applying this to '
    'each term of D(t) = 4t³ − 18t² + 24t + 90:'
)
add_bullet('d/dt(4t³) = 3 × 4 × t² = 12t²')
add_bullet('d/dt(−18t²) = 2 × (−18) × t = −36t')
add_bullet('d/dt(24t) = 1 × 24 × t⁰ = 24')
add_bullet('d/dt(90) = 0 (constant)')
add_para('Therefore: D\'(t) = 12t² − 36t + 24')
add_para('Differentiating again: D\'\'(t) = 24t − 36')

add_heading_styled('Step 2: Find Critical Points', level=3)
add_para(
    'Set D\'(t) = 0: 12t² − 36t + 24 = 0. Divide by 12: t² − 3t + 2 = 0. '
    'Factor: (t − 1)(t − 2) = 0. So critical points are at t = 1 and t = 2 hours.'
)

add_heading_styled('Step 3: Classify Using Second Derivative Test', level=3)
add_para('At t = 1: D\'\'(1) = 24(1) − 36 = −12. Since −12 < 0, this is a local maximum. '
         'D(1) = 4(1)³ − 18(1)² + 24(1) + 90 = 100 units.')
add_para('At t = 2: D\'\'(2) = 24(2) − 36 = 12. Since 12 > 0, this is a local minimum. '
         'D(2) = 4(2)³ − 18(2)² + 24(2) + 90 = 98 units.')

add_heading_styled('Step 4: Find Inflection Point', level=3)
add_para(
    'Set D\'\'(t) = 0: 24t − 36 = 0, so t = 1.5 hours. '
    'D(1.5) = 4(1.5)³ − 18(1.5)² + 24(1.5) + 90 = 99 units. '
    'This is where concavity changes from down to up.'
)

add_heading_styled('Step 5: Determine Intervals of Increase/Decrease', level=3)
add_para('Test points in each interval:')
add_bullet('t < 1 (e.g., t = 0): D\'(0) = 24 > 0 → Demand is INCREASING')
add_bullet('1 < t < 2 (e.g., t = 1.5): D\'(1.5) = −3 < 0 → Demand is DECREASING')
add_bullet('t > 2 (e.g., t = 3): D\'(3) = 24 > 0 → Demand is INCREASING')

# 2.4 Python Code - Line by Line
add_heading_styled('2.4 The Python Code - Line by Line', level=2)

add_para(
    'The script is in code/task2_calculus.py. Here is what each section does:'
)

add_para('Section 1 - Symbolic Setup (Lines 21-22):', bold=True)
add_para(
    'We use SymPy to define t as a symbolic variable and D(t) as a symbolic expression. This lets '
    'us compute derivatives exactly (not just numerically). SymPy manipulates mathematical symbols '
    'rather than numbers, so we get exact algebraic expressions for derivatives.'
)

add_para('Section 2 - Derivatives (Lines 34-55):', bold=True)
add_para(
    'sp.diff(D_sym, t_sym) computes the first derivative symbolically. We call sp.diff again on '
    'the result to get the second derivative. We also evaluate D, D\', and D"\ at several time '
    'points using direct numerical computation to see the pattern of values.'
)

add_para('Section 3 - Critical Points (Lines 61-87):', bold=True)
add_para(
    'We solve D\'(t) = 0 by hand (factoring the quadratic) and verify with sympy. Then we plug '
    'each critical point into D"\'(t) to classify it. D"\'(1) = −12 (negative, so local max). '
    'D"\'(2) = 12 (positive, so local min). The code prints the classification for each.'
)

add_para('Section 4 - Inflection Point (Lines 93-105):', bold=True)
add_para(
    'We set D\'\'(t) = 0 and solve: 24t − 36 = 0 gives t = 1.5. This is where the curve changes '
    'from bending downward to bending upward. The inflection point is significant because it signals '
    'the beginning of recovery before demand actually starts increasing.'
)

add_para('Section 5 - Visualisation (Lines 152-200):', bold=True)
add_para(
    'We create a two-panel figure. The top panel shows D(t) with the max (red), min (orange), '
    'and inflection point (gray dotted) marked. The bottom panel shows D\'(t) with green shading '
    'where demand is increasing (D\'(t) > 0) and red shading where it is decreasing (D\'(t) < 0). '
    'This dual-view approach helps connect the derivative behaviour to the original function.'
)

# 2.5 Understanding the Results
add_heading_styled('2.5 Understanding the Results', level=2)

add_para(
    'Demand starts at 90, rises to a morning peak of 100 at t = 1 hour, dips to 98 at t = 2 hours, '
    'then rises again to 154 by t = 4 hours. The inflection point at t = 1.5 is a subtle but important '
    'signal: even though demand is still decreasing at t = 1.5, the rate of decrease is slowing down - '
    'the system is preparing to recover.'
)

add_para(
    'Practical planning implications: the window between t = 1 and t = 2 hours is the best time '
    'for maintenance or resource reallocation because demand is at its lowest and the system is under '
    'least stress. After t = 2 hours, demand grows increasingly quickly, so planners need to ensure '
    'capacity is available for the sustained rise. The inflection point at t = 1.5 could serve as an '
    'early warning trigger for automated systems to begin ramping up resources before demand actually '
    'bottoms out at t = 2.'
)

# ============================================================
# TASK 3: FORECASTING ERROR RISK (PROBABILITY)
# ============================================================
add_heading_styled('Task 3: Forecasting Error Risk (Probability)', level=1)

# 3.1 What the Problem Is Asking
add_heading_styled('3.1 What the Problem Is Asking', level=2)

add_para(
    'We are creating daily forecasts of resource demand. These forecasts are never perfect - '
    'there will always be errors between predicted and actual demand. The errors come from two '
    'sources: behavioural variability (residents use more or less than expected on any given day) '
    'and environmental factors (weather, temperature, special events).'
)

add_para(
    'The key question for planners is: how big are the forecast errors likely to be? If errors are '
    'small, we can rely on the forecasts. If errors can be large, we need contingency buffers. '
    'We assume the errors follow a normal distribution with mean 0 (no systematic bias - forecasts '
    'are equally likely to overestimate as underestimate) and standard deviation 5 units.'
)

# 3.2 The Mathematics Behind It
add_heading_styled('3.2 The Mathematics Behind It - Normal Distribution', level=2)

add_para(
    'A normal distribution is a bell-shaped curve defined by two parameters: the mean (centre) and '
    'standard deviation (spread). For our errors, mean = 0 means the errors centre on zero (unbiased), '
    'and standard deviation = 5 means about 68% of errors are within ±5 units of zero.'
)

add_para(
    'The empirical rule (68-95-99.7 rule) is important:',
    bold=True
)
add_bullet('About 68% of values lie within ±1σ (i.e., ±5 units)')
add_bullet('About 95% of values lie within ±2σ (i.e., ±10 units)')
add_bullet('About 99.7% of values lie within ±3σ (i.e., ±15 units)')

add_para(
    'To compute exact probabilities, we use z-scores. A z-score measures how many standard deviations '
    'a particular threshold is from the mean: z = (x − μ)/σ. For threshold 10: z = (10 − 0)/5 = 2.0. '
    'We then use the cumulative distribution function (CDF) to find the probability that a value '
    'exceeds this threshold.'
)

# 3.3 Step-by-Step Solution
add_heading_styled('3.3 Step-by-Step Solution', level=2)

add_heading_styled('Step 1: Computing Tail Probabilities', level=3)
add_para(
    'For a threshold t, we want P(|error| > t), which is the probability the error magnitude exceeds t. '
    'Since the normal distribution is symmetric, P(|error| > t) = 2 × P(error > t).'
)

add_para(
    'P(error > t) is found using the survival function = 1 − CDF(z) where z = t/5.',
    bold=True
)
add_bullet('For t = 5: z = 1.0, P(|error| > 5) = 2 × (1 − 0.8413) = 0.3173 (31.73%)')
add_bullet('For t = 10: z = 2.0, P(|error| > 10) = 2 × (1 − 0.9772) = 0.0456 (4.56%)')
add_bullet('For t = 15: z = 3.0, P(|error| > 15) = 2 × (1 − 0.9987) = 0.0027 (0.27%)')

add_heading_styled('Step 2: Buffer Coverage Interpretation', level=3)
add_para(
    'The complement probabilities give us the buffer coverage:'
)
add_bullet('Buffer ±5: covers 68.27% of days → 31.73% of days will exceed this')
add_bullet('Buffer ±10: covers 95.45% of days → only 4.55% of days exceed')
add_bullet('Buffer ±15: covers 99.73% of days → only 0.27% exceed (about 1 day per year)')

# 3.4 Python Code - Line by Line
add_heading_styled('3.4 The Python Code - Line by Line', level=2)

add_para('The script is in code/task3_probability.py. Here is what each section does:')

add_para('Section 1 - Distribution Parameters (Lines 18-19):', bold=True)
add_para(
    'We set mu = 0 (mean) and sigma = 5 (standard deviation). These define the normal distribution '
    'that models the forecast errors.'
)

add_para('Section 2 - Probability Calculations (Lines 24-35):', bold=True)
add_para(
    'We define thresholds of interest: 5, 10, and 15 units (corresponding to 1, 2, and 3 standard '
    'deviations). For each threshold, we compute P(|error| > t) using stats.norm.cdf(t) which gives '
    'the cumulative probability up to t. Since we want the tail probability beyond t, we compute '
    '1 - cdf(t), and since we want both tails (positive and negative errors), we multiply by 2.'
)

add_para('Section 3 - Visualisation (Lines 43-80):', bold=True)
add_para(
    'We create a normal curve plot using matplotlib. The x-axis runs from -4σ to +4σ to show the '
    'full range of possible errors. The main curve is plotted using stats.norm.pdf. We then fill '
    'the tail regions beyond ±5, ±10, and ±15 with different colours and transparency. Vertical '
    'dashed lines mark each threshold. Annotations on the plot show the probability of exceeding '
    'each threshold. The figure is saved as fig_task3_normal_curve.png.'
)

add_para('Section 4 - Risk Management Interpretation (Lines 83-87):', bold=True)
add_para(
    'We print a summary table showing what buffer level covers what percentage of days. This '
    'directly informs contingency planning: a ±10 unit buffer will be adequate for 95% of days, '
    'while a ±15 unit buffer covers 99.7% of days but ties up more resources than necessary.'
)

# 3.5 Understanding the Results
add_heading_styled('3.5 Understanding the Results', level=2)

add_para(
    'The results show that forecast errors follow a predictable pattern described by the normal '
    'distribution. About two-thirds of the time (68.3%), the error will be within ±5 units. '
    'About 95% of the time, the error will be within ±10 units. Only about once in every 370 days '
    '(0.27%) will the error exceed ±15 units.'
)

add_para(
    'For practical planning, a contingency buffer of ±10 units provides a good balance: it covers '
    'almost all normal operations (95.4% of days) while not wasting resources that could be used '
    'elsewhere. A ±15 unit buffer is overly conservative - covering 99.7% of days but requiring '
    'three times the reserve capacity. The remaining 0.3% of exceptional days (about one per year) '
    'can be handled through emergency procedures rather than standing reserves.'
)

add_para(
    'It is important to note the limitations: the normal distribution assumption means we are '
    'assuming errors are symmetric and not too extreme. In practice, during unusual events '
    '(heatwaves, festivals, major disruptions), errors may be larger. The standard deviation '
    'σ = 5 should be periodically recalibrated using actual forecast performance data.'
)

# ============================================================
# TASK 4: CONSUMPTION BY BUILDING TYPE (STATISTICS)
# ============================================================
add_heading_styled('Task 4: Consumption by Building Type (Statistics)', level=1)

# 4.1 What the Problem Is Asking
add_heading_styled('4.1 What the Problem Is Asking', level=2)

add_para(
    'We have data on daily water consumption from three types of residential buildings: '
    'Apartments (8 observations), Terraced houses (10 observations), and Detached houses '
    '(6 observations). The sample sizes are unequal because data collection in real-world '
    'settings rarely produces perfectly balanced groups. The question is: do these building '
    'types have different average daily consumption, or could the observed differences simply '
    'be due to random chance?'
)

add_para('The data:')
add_bullet('Apartments (n=8): 16.8, 17.2, 16.5, 17.9, 18.1, 17.4, 16.9, 17.6')
add_bullet('Terraced (n=10): 18.4, 18.9, 19.1, 18.7, 19.3, 18.6, 19.0, 18.8, 19.2, 18.5')
add_bullet('Detached (n=6): 20.1, 19.7, 20.4, 19.9, 20.6, 20.2')

# 4.2 The Mathematics Behind It
add_heading_styled('4.2 The Mathematics Behind It - One-Way ANOVA', level=2)

add_para(
    'One-way ANOVA (Analysis of Variance) is a statistical method used to compare means '
    'across three or more groups. It is called "one-way" because there is one categorical '
    'independent variable (building type) and one continuous dependent variable (consumption).'
)

add_para('Key concepts in ANOVA:', bold=True)

add_para('Null Hypothesis (H0):', bold=True)
add_para('All group means are equal: mean_apartments = mean_terraced = mean_detached. '
         'Any observed differences are due to random variation within the groups.')

add_para('Alternative Hypothesis (H1):', bold=True)
add_para('At least one group mean differs from the others. The observed differences are '
         'too large to be explained by random variation alone.')

add_para('The F-statistic:', bold=True)
add_para(
    'ANOVA works by comparing two sources of variation: (1) variation between groups '
    '(how far each group mean is from the overall mean) and (2) variation within groups '
    '(how far individual observations are from their own group mean). The F-statistic is '
    'the ratio of between-group variation to within-group variation. A large F-value means '
    'the differences between groups are large compared to the natural variation within groups, '
    'suggesting a real difference exists.'
)

add_para('Assumptions:', bold=True)
add_para(
    'Three key assumptions must be checked: (1) Independence - observations are independent '
    'of each other (satisfied by the study design), (2) Normality - each group follows a '
    'normal distribution (checked with Shapiro-Wilk test), and (3) Homogeneity of variances '
    '- the variance is approximately equal across groups (checked with Levene\'s test). '
    'ANOVA is robust to mild violations of normality, especially with equal or nearly equal '
    'sample sizes.'
)

add_para('Tukey HSD (Honestly Significant Difference):', bold=True)
add_para(
    'If ANOVA finds a significant difference, we need to know which pairs of groups differ. '
    'Tukey\'s HSD test performs all pairwise comparisons while controlling the "family-wise '
    'error rate" - the probability of making at least one false discovery across all comparisons. '
    'It handles unequal sample sizes automatically.'
)

# 4.3 Step-by-Step Solution
add_heading_styled('4.3 Step-by-Step Solution', level=2)

add_heading_styled('Step 1: Compute Descriptive Statistics', level=3)
add_para(
    'For each group, we calculate the mean, standard deviation, minimum, and maximum. '
    'This gives us a first look at the data:'
)
add_bullet('Apartments: mean = 17.30, SD = 0.56, range [16.5, 18.1]')
add_bullet('Terraced: mean = 18.85, SD = 0.30, range [18.4, 19.3]')
add_bullet('Detached: mean = 20.15, SD = 0.33, range [19.7, 20.6]')
add_para(
    'Already we can see clear separation: the means are approximately 17.3, 18.9, and 20.2, '
    'with small standard deviations relative to the gaps between groups.'
)

add_heading_styled('Step 2: Check Assumptions', level=3)
add_para('Normality (Shapiro-Wilk Test):', bold=True)
add_para(
    'For each group, we compute the Shapiro-Wilk W statistic and its p-value. A p-value > 0.05 '
    'means we cannot reject normality. All three groups pass: Apartments (W=0.97, p=0.92), '
    'Terraced (W=0.97, p=0.89), Detached (W=0.99, p=0.99). This confirms the normality '
    'assumption is reasonable.'
)

add_para('Equal Variances (Levene\'s Test):', bold=True)
add_para(
    'Levene\'s test returns Levene = 2.49, p = 0.1071. Since p > 0.05, the variances are '
    'not significantly different, so the equal variance assumption is met. Standard one-way '
    'ANOVA is appropriate.'
)

add_heading_styled('Step 3: Perform One-Way ANOVA', level=3)
add_para(
    'scipy.stats.f_oneway computes the F-statistic and p-value. The F-statistic is calculated '
    'by computing the between-group mean square and dividing by the within-group mean square. '
    'For our data: F(2, 21) = 84.95, p < 0.000001. This is a very large F-value, indicating '
    'that the between-group variation dwarfs the within-group variation.'
)

add_para(
    'Since p < 0.05, we reject the null hypothesis. There is strong statistical evidence that '
    'at least one building type has a different mean consumption from the others.'
)

add_heading_styled('Step 4: Post-Hoc Testing (Tukey HSD)', level=3)
add_para(
    'The Tukey HSD test from statsmodels performs all three pairwise comparisons:'
)
add_bullet('Apartments vs Detached: difference = 2.85, p < 0.001, 95% CI [2.29, 3.41]')
add_bullet('Apartments vs Terraced: difference = 1.55, p < 0.001, 95% CI [1.06, 2.04]')
add_bullet('Detached vs Terraced: difference = -1.30, p < 0.001, 95% CI [-1.83, -0.77]')
add_para(
    'All three comparisons are statistically significant (p < 0.05). The confidence intervals '
    'do not include zero, confirming that each building type has a distinct consumption level. '
    'The "reject=True" column in the output confirms rejection of the null hypothesis for '
    'each pair.'
)

# 4.4 Python Code - Line by Line
add_heading_styled('4.4 The Python Code - Line by Line', level=2)

add_para('The script is in code/task4_statistics.py. Here is what each section does:')

add_para('Section 1 - Data Setup (Lines 16-26):', bold=True)
add_para(
    'We define three NumPy arrays containing the consumption data for each building type. '
    'The groups dictionary maps group names to their data arrays for easy iteration. '
    'The sample sizes are explicitly n=8, n=10, and n=6.'
)

add_para('Section 2 - Descriptive Statistics (Lines 32-42):', bold=True)
add_para(
    'We print a formatted table showing n, mean, standard deviation (using ddof=1 for sample '
    'SD), minimum, and maximum for each group. This provides an initial overview before any '
    'inferential testing.'
)

add_para('Section 3 - Assumption Checks (Lines 48-80):', bold=True)
add_para(
    'scipy.stats.shapiro computes the Shapiro-Wilk test for each group. The null hypothesis '
    'is that the data follows a normal distribution. A high p-value (> 0.05) means we cannot '
    'reject normality. scipy.stats.levene tests whether group variances are equal. Its null '
    'hypothesis is that all variances are equal. A p-value > 0.05 means the equal variance '
    'assumption is met.'
)

add_para('Section 4 - ANOVA (Lines 86-103):', bold=True)
add_para(
    'scipy.stats.f_oneway performs the one-way ANOVA. It takes the data arrays as separate '
    'arguments and returns the F-statistic and p-value. We print the hypotheses, significance '
    'level, and decision rule. If p < 0.05, we reject H0 and conclude at least one group differs.'
)

add_para('Section 5 - Tukey HSD (Lines 113-127):', bold=True)
add_para(
    'We concatenate all data into a single array and create a labels array indicating which '
    'group each value belongs to. pairwise_tukeyhsd from statsmodels performs the post-hoc '
    'test. It prints a table showing each pair, the mean difference, adjusted p-value, '
    'confidence interval, and whether to reject the null hypothesis for that pair.'
)

add_para('Section 6 - Visualisation (Lines 133-165):', bold=True)
add_para(
    'We use seaborn boxplot to show the distribution of each group, with stripplot overlay '
    'showing individual data points. Red diamond markers indicate group means. The figure '
    'is saved as fig_task4_boxplot.png. This visualisation helps stakeholders see the clear '
    'separation between groups at a glance.'
)

# 4.5 Understanding the Results
add_heading_styled('4.5 Understanding the Results', level=2)

add_para(
    'The analysis conclusively shows that different building types have different consumption '
    'patterns. The ANOVA F(2,21) = 84.95, p < 0.001 is a very strong result - the probability '
    'of observing such extreme differences by chance alone is less than 1 in a million. All '
    'three pairwise comparisons (via Tukey HSD) are significant, meaning each building type '
    'is distinct from the others.'
)

add_para(
    'The practical significance is also substantial. Detached homes consume 2.85 units more '
    'per day than apartments, a 16.5% increase. For a city with thousands of households, this '
    'represents a massive cumulative difference. Energy efficiency programmes should prioritise '
    'detached homes, where the potential savings per household are greatest.'
)

add_para(
    'The lower consumption in apartments likely reflects their shared infrastructure, smaller '
    'living spaces, and greater thermal efficiency from having neighbouring units. Terraced '
    'houses fall in between, benefiting from shared walls but having more space than '
    'apartments. These results justify targeted, building-type-specific intervention strategies.'
)

# ============================================================
# SAVE
# ============================================================
output_path = 'code_explanation.docx'
doc.save(output_path)
print(f'Code explanation saved to {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')
